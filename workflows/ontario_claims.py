# -*- coding: utf-8 -*-
"""Generate Ontario Schedule A and Plaintiff's Claim from source files."""

from __future__ import annotations

import io
import re
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Dict, Iterator, Optional, Tuple

import pdfplumber
from docx import Document

COMPANY_KEYWORDS = (
    " INC",
    " LTD",
    " LLC",
    " CO ",
    " COMPANY",
    " CORP",
    " CORP.",
    " BANK",
    " LIMITED",
    " PLC",
    " LLP",
)

PROVINCE_NAMES = {
    "ON": "Ontario",
    "QC": "Quebec",
    "NS": "Nova Scotia",
    "NB": "New Brunswick",
    "MB": "Manitoba",
    "BC": "British Columbia",
    "PE": "Prince Edward Island",
    "SK": "Saskatchewan",
    "AB": "Alberta",
    "NL": "Newfoundland and Labrador",
    "NT": "Northwest Territories",
    "YT": "Yukon",
    "NU": "Nunavut",
}

SCHEDULE_PLACEHOLDERS = (
    "SIMPLE NAME INSERT",
    "DATE OF LAST CHARGE INSERT",
    "DATEOFLASTPAYMENTINSERT",
    "DEMAND LETTER DATE INSERT",
    "TOTALDEBTINSERT",
)
CLAIM_PLACEHOLDERS = (
    "LAST NAME INSERT",
    "FIRST NAME INSERT",
    "AKA INSERT",
    "PHONENOINSERT",
    "TODAYS DATE",
    "YEARENDINSERT",
)

MODULE_DIR = Path(__file__).resolve().parent
ASSET_DIR = MODULE_DIR.parent / "assets" / "ontario"
DEFAULT_SCHEDULE_TEMPLATE = ASSET_DIR / "Schedule A - AMEX ON.docx"
DEFAULT_CLAIM_TEMPLATE = ASSET_DIR / "Plaintiffs Claim Form 7A.docx"


class OntarioClaimsError(RuntimeError):
    """Raised when the Ontario claims workflow cannot complete."""


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def to_title_case(text: str) -> str:
    if not text:
        return text
    return text.title() if text.isupper() else text


def format_currency(value: str | Decimal | float) -> str:
    if isinstance(value, Decimal):
        amount = value
    else:
        cleaned = str(value).replace("$", "").replace(",", "").strip()
        amount = Decimal(cleaned or "0")
    amount = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"${amount:,.2f}"


def canonicalize_postal_code(code: str) -> str:
    cleaned = re.sub(r"[^A-Z0-9]", "", code.upper())
    if len(cleaned) == 6:
        return f"{cleaned[:3]} {cleaned[3:]}"
    return cleaned


def parse_compact_date(token: str) -> datetime:
    cleaned = token.replace(" ", "")
    return datetime.strptime(cleaned, "%b%d,%Y")


def format_short_month(date_obj: datetime) -> str:
    abbr = date_obj.strftime("%b")
    if abbr not in {"May", "Jun", "Jul"} and not abbr.endswith("."):
        abbr = f"{abbr}."
    return f"{abbr} {date_obj.strftime('%d, %Y')}"


def format_full_date(date_obj: datetime) -> str:
    return date_obj.strftime("%B %d, %Y").replace(" 0", " ")


def format_claim_date(date_obj: datetime) -> str:
    return date_obj.strftime("%B %d").replace(" 0", " ")


def format_phone(number: str) -> str:
    digits = re.sub(r"\D", "", number)
    if len(digits) == 10:
        return f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
    return digits


def parse_user_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    value = value.strip()
    formats = [
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%b %d, %Y",
        "%B %d, %Y",
        "%d %b %Y",
        "%d %B %Y",
    ]
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    # As a fallback, return None so the caller can leave the field blank
    return None


def load_pdf_text_and_lines(path: Path) -> Tuple[str, list[str]]:
    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    text = "\n".join(pages)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return text, lines


def resolve_statement_date(token: str, start: datetime, end: datetime) -> datetime:
    match = re.match(r"([A-Za-z]{3})(\d{1,2})", token)
    if not match:
        raise OntarioClaimsError(f"Unexpected date token: {token}")
    month_name, day = match.groups()
    month = datetime.strptime(month_name, "%b").month
    day_int = int(day)
    year = start.year
    if end.year > start.year and month <= end.month and month < start.month:
        year = end.year
    return datetime(year, month, day_int)


def split_name_for_claim(full_name: str) -> Tuple[str, str, str]:
    upper = f" {full_name.upper()} "
    if any(keyword in upper for keyword in COMPANY_KEYWORDS):
        return full_name, "", ""
    parts = full_name.split()
    if not parts:
        return "", "", ""
    if len(parts) == 1:
        return parts[0], "", ""
    first = parts[0]
    last = parts[-1]
    middle = " ".join(parts[1:-1])
    return last, first, middle


def build_alias(first: str, middle: str, last: str) -> str:
    if not middle:
        return ""
    tokens = middle.split()
    formatted_parts = []
    for token in tokens:
        stripped = token.replace(".", "")
        if len(stripped) == 1:
            formatted_parts.append(f"{stripped}.")
        else:
            formatted_parts.append(stripped)
    alias = normalize_whitespace(f"{first} {' '.join(formatted_parts)} {last}")
    full = normalize_whitespace(f"{first} {middle} {last}")
    if alias.upper() == full.upper():
        return ""
    return alias.upper()


def sanitize_for_filename(text: str, space_replacement: str = " ") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 _-]", "", text).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    if space_replacement != " ":
        cleaned = cleaned.replace(" ", space_replacement)
    return cleaned or "output"


def identify_template(path: Path) -> Optional[str]:
    document = Document(path)
    text_chunks = []
    for paragraph in document.paragraphs:
        text_chunks.append(paragraph.text or "")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text_chunks.append(paragraph.text or "")
    full_text = "\n".join(text_chunks)
    schedule_score = sum(1 for key in SCHEDULE_PLACEHOLDERS if key in full_text)
    claim_score = sum(1 for key in CLAIM_PLACEHOLDERS if key in full_text)
    if schedule_score == 0 and claim_score == 0:
        return None
    return "SCHEDULE" if schedule_score >= claim_score else "CLAIM"


def parse_mrs_statement(path: Path) -> Dict[str, object]:
    text, lines = load_pdf_text_and_lines(path)
    header_match = re.search(
        r"Prepared For Account Number Opening Date Closing Date\n([A-Z ,]+?)\s+X{4}",
        text,
    )
    if not header_match:
        raise OntarioClaimsError("Could not locate account header in the MRS statement.")
    full_name = normalize_whitespace(header_match.group(1))
    header_line = next(
        line for line in lines if line.startswith(full_name) and "XXXX" in line
    )
    card_number_match = re.search(r"(X{3,}[ X\d-]*\d{4,})", header_line)
    if not card_number_match:
        raise OntarioClaimsError(
            f"Unable to parse the masked card number from: {header_line}"
        )
    card_number = card_number_match.group(1)
    suffix_match = re.search(r"(\d{4})\s*$", card_number)
    card_suffix = suffix_match.group(1) if suffix_match else card_number[-4:]
    date_tokens = re.findall(r"[A-Za-z]{3}\d{2}, \d{4}", header_line)
    if len(date_tokens) != 2:
        raise OntarioClaimsError("Could not determine the statement period from MRS.")
    start_date = parse_compact_date(date_tokens[0])
    end_date = parse_compact_date(date_tokens[1])
    card_type_candidates = [
        line for line in lines[:40]
        if "card" in line.lower() and "american express" in line.lower()
    ]
    card_type = card_type_candidates[0] if card_type_candidates else "American Express Credit Card"
    try:
        name_index = lines.index(full_name)
    except ValueError as exc:
        raise OntarioClaimsError("Unable to locate mailing block name in MRS.") from exc
    address_lines = []
    for line in lines[name_index + 1:]:
        if not re.fullmatch(r"[A-Z0-9 ,#'&.-]+", line):
            break
        address_lines.append(line)
    street_parts = address_lines[:-1] if len(address_lines) > 1 else address_lines
    city_line = address_lines[-1] if address_lines else ""
    city_tokens = city_line.split()
    city_raw = normalize_whitespace(" ".join(city_tokens[:-2])) if len(city_tokens) >= 3 else ""
    province = city_tokens[-2] if len(city_tokens) >= 2 else ""
    postal = canonicalize_postal_code(city_tokens[-1]) if city_tokens else ""
    schedule_address_parts = list(street_parts)
    if city_raw:
        schedule_address_parts.append(f"{city_raw} {province} {postal}".strip())
    schedule_address = ", ".join(part for part in schedule_address_parts if part)
    street_unit = ", ".join(street_parts).strip(", ")
    summary_patterns = {
        "previous_balance": r"Previous Balance \$([0-9,]+\.\d{2})",
        "less_payments": r"Less Payments \$([0-9,]+\.\d{2})",
        "plus_interest": r"Plus Interest \$([0-9,]+\.\d{2})",
        "plus_fees": r"Plus Fees \$([0-9,]+\.\d{2})",
        "new_balance": r"Equals New Balance \$([0-9,]+\.\d{2})",
    }
    summary: Dict[str, Decimal] = {}
    for key, pattern in summary_patterns.items():
        match = re.search(pattern, text)
        if not match:
            raise OntarioClaimsError(f"Missing '{key}' in the account summary.")
        summary[key] = Decimal(match.group(1).replace(",", ""))
    province_full = PROVINCE_NAMES.get(province.upper(), province.upper())
    interest_rate = ""
    for idx, line in enumerate(lines):
        if line.strip().lower() == "purchases":
            for lookahead in range(1, 4):
                if idx + lookahead >= len(lines):
                    break
                percents = re.findall(r"\d+\.\d+%", lines[idx + lookahead])
                if percents:
                    interest_rate = max(
                        percents, key=lambda item: Decimal(item.rstrip("%"))
                    )
                    break
            if interest_rate:
                break
    if not interest_rate:
        raise OntarioClaimsError("Could not locate the purchases interest rate in MRS.")
    return {
        "full_name": full_name,
        "card_number": card_number,
        "card_suffix": card_suffix,
        "card_type": card_type,
        "start_date": start_date,
        "end_date": end_date,
        "date_range": f"{format_short_month(start_date)} – {format_short_month(end_date)}",
        "schedule_address": schedule_address,
        "street_unit": street_unit,
        "city": city_raw.upper() if city_raw else "",
        "city_title": to_title_case(city_raw),
        "province": province.upper(),
        "province_full": province_full,
        "postal": postal,
        "summary": summary,
        "interest_rate": interest_rate,
    }


def parse_last_purchase_date(path: Path) -> datetime:
    text, lines = load_pdf_text_and_lines(path)
    header_match = re.search(
        r"Prepared For Account Number Opening Date Closing Date\n([A-Z ,]+?)\s+X{4}",
        text,
    )
    if not header_match:
        raise OntarioClaimsError("Could not locate header in MRC statement.")
    full_name = normalize_whitespace(header_match.group(1))
    header_line = next(
        line for line in lines if line.startswith(full_name) and "XXXX" in line
    )
    date_tokens = re.findall(r"[A-Za-z]{3}\d{2}, \d{4}", header_line)
    if len(date_tokens) != 2:
        raise OntarioClaimsError("Unable to parse date range in MRC statement.")
    start_date = parse_compact_date(date_tokens[0])
    end_date = parse_compact_date(date_tokens[1])
    in_transactions = False
    candidate_dates: list[datetime] = []
    for line in lines:
        if line.startswith("New Transactions for"):
            in_transactions = True
            continue
        if in_transactions and (
            line.startswith("Total of New Transactions")
            or line.startswith("Other Account Transactions")
        ):
            break
        if not in_transactions:
            continue
        tokens = line.split()
        if len(tokens) < 3 or not re.match(r"^[A-Za-z]{3}\d{1,2}$", tokens[0]):
            continue
        description = " ".join(tokens[2:-1]).upper()
        if any(
            bad in description
            for bad in ("PAYMENT", "INTEREST", "FEE", "ADJUSTMENT", "CREDIT", "RETURN")
        ):
            continue
        amount_token = tokens[-1].replace("$", "").replace(",", "")
        if amount_token.endswith("CR"):
            continue
        sign = 1
        stripped = amount_token.strip()
        if stripped.startswith("(") and stripped.endswith(")"):
            sign = -1
            stripped = stripped[1:-1]
        if stripped.startswith("-"):
            sign = -1
            stripped = stripped[1:]
        try:
            amount_val = Decimal(stripped)
        except Exception:
            amount_val = Decimal("0")
        if sign < 0 or amount_val <= 0:
            continue
        date_obj = resolve_statement_date(tokens[0], start_date, end_date)
        candidate_dates.append(date_obj)
    if not candidate_dates:
        raise OntarioClaimsError("No purchase transactions were detected in MRC.")
    return max(candidate_dates)


def parse_last_payment_date(path: Path) -> datetime:
    text, lines = load_pdf_text_and_lines(path)
    header_match = re.search(
        r"Prepared For Account Number Opening Date Closing Date\n([A-Z ,]+?)\s+X{4}",
        text,
    )
    if not header_match:
        raise OntarioClaimsError("Could not locate header in MRP statement.")
    full_name = normalize_whitespace(header_match.group(1))
    header_line = next(
        line for line in lines if line.startswith(full_name) and "XXXX" in line
    )
    date_tokens = re.findall(r"[A-Za-z]{3}\d{2}, \d{4}", header_line)
    if len(date_tokens) != 2:
        raise OntarioClaimsError("Unable to parse date range in MRP statement.")
    start_date = parse_compact_date(date_tokens[0])
    end_date = parse_compact_date(date_tokens[1])
    payment_dates: list[datetime] = []
    for line in lines:
        if "PAYMENT RECEIVED - THANK YOU" not in line.upper():
            continue
        tokens = line.split()
        if not tokens:
            continue
        date_obj = resolve_statement_date(tokens[0], start_date, end_date)
        payment_dates.append(date_obj)
    if not payment_dates:
        raise OntarioClaimsError("No payments were detected in MRP.")
    return max(payment_dates)


def extract_credit_report_data(path: Path) -> Dict[str, str]:
    text, lines = load_pdf_text_and_lines(path)
    subject_line = next((line for line in lines if line.startswith("Subject ")), "")
    if not subject_line:
        raise OntarioClaimsError("Subject line missing in credit report.")
    subject_tokens = subject_line.split()[1:]
    birth_token = subject_tokens[-1]
    name_tokens = subject_tokens[:-1]
    if not name_tokens:
        raise OntarioClaimsError("Unable to parse name from credit report.")
    last_name = name_tokens[0]
    given_names = " ".join(token.replace("/", " ") for token in name_tokens[1:])
    full_name = normalize_whitespace(f"{given_names} {last_name}")
    header = "On File Last Inq Current Residence Telephone Prev phone"
    if header not in lines:
        raise OntarioClaimsError("Telephone header missing in credit report.")
    idx = lines.index(header)
    data_tokens = lines[idx + 1].split()
    if len(data_tokens) < 4:
        raise OntarioClaimsError("Telephone row malformed in credit report.")
    current_phone = format_phone(data_tokens[2])
    prev_phone = format_phone(data_tokens[3]) if len(data_tokens) > 3 else ""
    phone = current_phone or prev_phone
    return {"full_name": full_name, "phone": phone, "birth": birth_token}


def apply_mapping(text: str, mapping: Dict[str, str]) -> str:
    result = text
    for key in sorted(mapping, key=len, reverse=True):
        result = result.replace(key, mapping[key])
    return result


def iter_paragraphs(element) -> Iterator:
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            yield paragraph
    if hasattr(element, "tables"):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return
        original = "".join(run.text for run in paragraph.runs)
        updated = apply_mapping(original, mapping)
        if updated != original:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""

    for paragraph in iter_paragraphs(doc):
        replace_in_paragraph(paragraph)

    for section in doc.sections:
        for header in (section.header, section.footer):
            for paragraph in iter_paragraphs(header):
                replace_in_paragraph(paragraph)


def build_schedule_mapping(
    mrs_data: Dict[str, object],
    last_charge: datetime,
    last_payment: datetime,
    alias: str,
    demand_letter_date: str,
) -> Dict[str, str]:
    summary: Dict[str, Decimal] = mrs_data["summary"]  # type: ignore[assignment]
    simple_name = mrs_data["full_name"]
    if alias:
        simple_name = f"{simple_name} a.k.a. {alias}"
    return {
        "SIMPLE NAME INSERT": simple_name,
        "ADRESS INSERT": f"{mrs_data['city_title']}, {mrs_data['province_full']}",
        "CARDNUMBERINSERT": mrs_data["card_suffix"],
        "CARD TYPE INSERT": mrs_data["card_type"],
        "DATE OF LAST CHARGE INSERT": format_full_date(last_charge),
        "DATEOFLASTPAYMENTINSERT": format_full_date(last_payment),
        "INTEREST RATE INSERT": mrs_data["interest_rate"],
        "DATE RANGE INSERT": mrs_data["date_range"],
        "DEBT INSERT": format_currency(summary["previous_balance"]),
        "ACCRUED INTEREST INSERT": format_currency(summary["plus_interest"]),
        "FINSERT": format_currency(summary["plus_fees"]),
        "PAYMENTS INSERT": format_currency(summary["less_payments"]),
        "TOTALDEBTINSERT": format_currency(summary["new_balance"]),
        "DEMAND LETTER DATE INSERT": demand_letter_date,
        "Ottawa Small Claims Court": f"{mrs_data['city_title']} Small Claims Court",
    }


def build_claim_mapping(
    mrs_data: Dict[str, object],
    phone: str,
    prepared_date: datetime,
    name_parts: Tuple[str, str, str],
    alias: str,
) -> Dict[str, str]:
    summary: Dict[str, Decimal] = mrs_data["summary"]  # type: ignore[assignment]
    last_name, first_name, middle_names = name_parts
    aka_value = alias
    return {
        "LAST NAME INSERT": last_name,
        "FIRST NAME INSERT": first_name,
        "AKA INSERT": aka_value,
        "ADRESS INSERT": mrs_data["street_unit"],
        "CITY INSERT": mrs_data["city"],
        "PROVINSERT": mrs_data["province"],
        "POSTALCODEINSERT": mrs_data["postal"],
        "PHONENOINSERT": phone,
        "DEBTINSERT": format_currency(summary["new_balance"]),
        "TODAYS DATE": format_claim_date(prepared_date),
        "YEARENDINSERT": prepared_date.strftime("%y"),
    }


def document_to_bytes(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_claim_documents(
    *,
    mrs_path: Path,
    mrc_path: Path,
    mrp_path: Path,
    cbr_path: Path,
    schedule_template: Path = DEFAULT_SCHEDULE_TEMPLATE,
    claim_template: Path = DEFAULT_CLAIM_TEMPLATE,
    demand_letter_date: Optional[str] = None,
    claim_prepared_date: Optional[str] = None,
) -> Dict[str, Tuple[str, bytes]]:
    if not schedule_template.exists():
        raise OntarioClaimsError(f"Schedule template not found at {schedule_template}")
    if not claim_template.exists():
        raise OntarioClaimsError(f"Claim template not found at {claim_template}")
    mrs_data = parse_mrs_statement(mrs_path)
    last_charge = parse_last_purchase_date(mrc_path)
    last_payment = parse_last_payment_date(mrp_path)
    credit_data = extract_credit_report_data(cbr_path)
    parsed_demand = parse_user_date(demand_letter_date)
    demand_date_str = format_full_date(parsed_demand) if parsed_demand else (demand_letter_date or "")
    prepared_dt = parse_user_date(claim_prepared_date) or datetime.today()
    name_parts = split_name_for_claim(credit_data["full_name"])
    alias = build_alias(name_parts[1], name_parts[2], name_parts[0])
    schedule_mapping = build_schedule_mapping(
        mrs_data,
        last_charge,
        last_payment,
        alias,
        demand_date_str,
    )
    claim_mapping = build_claim_mapping(
        mrs_data,
        credit_data["phone"],
        prepared_dt,
        name_parts,
        alias,
    )
    schedule_doc = Document(schedule_template)
    replace_placeholders(schedule_doc, schedule_mapping)
    schedule_bytes = document_to_bytes(schedule_doc)
    claim_doc = Document(claim_template)
    replace_placeholders(claim_doc, claim_mapping)
    claim_bytes = document_to_bytes(claim_doc)
    schedule_filename = (
        f"filled_template_Schedule A - 1 Credit Card_{sanitize_for_filename(mrs_data['full_name'])}.docx"
    )
    last_name, first_name, _ = name_parts
    claim_suffix_parts = [sanitize_for_filename(last_name, "_")]
    if first_name:
        claim_suffix_parts.append(sanitize_for_filename(first_name, "_"))
    claim_filename = f"Plaintiffs_Claim_{'_'.join(claim_suffix_parts)}.docx"
    return {
        "schedule": (schedule_filename, schedule_bytes),
        "claim": (claim_filename, claim_bytes),
    }

# -*- coding: utf-8 -*-
"""Generate BC Claims Schedule A and Notice of Claim from four source PDFs."""

from __future__ import annotations

import argparse
import io
import re
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import pdfplumber
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject

MODULE_DIR = Path(__file__).resolve().parent
INPUT_DIR = MODULE_DIR / "input"
OUTPUT_DIR = MODULE_DIR.parent / "output"
ASSET_DIR = MODULE_DIR.parent / "assets" / "claims"
SCHEDULE_TEMPLATE = ASSET_DIR / "Schedule A - 1 Credit Card.docx"
NOTICE_TEMPLATE = ASSET_DIR / "Notice of Claim_tofill.pdf"

FILE_TYPE_PATTERNS = {
    "MRP": ("*MRP*.pdf",),
    "MRC": ("*MRC*.pdf",),
    "MRS": ("*MRS*.pdf",),
    "CBR": ("*CBR*.pdf", "*Credit Report*.pdf"),
}


# ---------------------------------------------------------------------------
# Utility helpers

def normalize_unicode(value: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\ufffd": "",
    }
    for raw, replacement in replacements.items():
        value = value.replace(raw, replacement)
    return value


def safe_filename(value: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", value)
    sanitized = re.sub(r"_{2,}", "_", sanitized).strip(" .")
    return sanitized or "output"


def auto_pick_file(directory: Path, patterns: Iterable[str]) -> Optional[Path]:
    for pattern in patterns:
        matches = sorted(directory.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True)
        if matches:
            return matches[0]
    return None


def to_money_str(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return ""
        cleaned = stripped.replace("$", "").replace(",", "")
        try:
            dec = Decimal(cleaned)
            return f"${dec:,.2f}"
        except InvalidOperation:
            return stripped
    try:
        dec = Decimal(str(value))
        return f"${dec:,.2f}"
    except InvalidOperation:
        return str(value)


def normalize_phone(value: str) -> str:
    digits = re.sub(r"\D+", "", value or "")
    return digits if len(digits) >= 10 else ""


def parse_bc_address(addr: str) -> Tuple[str, str, str, str]:
    s = (addr or "").strip()
    if not s:
        return "", "", "", ""
    match = re.search(
        r"^(.*?),\s*([^,]+?)\s+([A-Za-z]{2})\s+([A-Za-z]\d[A-Za-z]\s?\d[A-Za-z]\d)$",
        s,
    )
    if match:
        street, city, prov, postal = match.groups()
        return (
            street.strip(),
            city.strip(),
            prov.strip().upper(),
            postal.replace(" ", "").upper(),
        )
    street, city, prov, postal = s, "", "", ""
    postal_match = re.search(r"([A-Za-z]\d[A-Za-z]\s?\d[A-Za-z]\d)$", s)
    if postal_match:
        postal = postal_match.group(1).replace(" ", "").upper()
        street = s.replace(postal_match.group(1), "").strip(", ")
    prov_match = re.search(r"\b(AB|BC|MB|NB|NL|NS|NT|NU|ON|PE|QC|SK|YT)\b", s)
    if prov_match:
        prov = prov_match.group(1)
        parts = s.split(",")
        if len(parts) >= 2:
            possible_city = parts[1].strip().split()
            city = " ".join(t for t in possible_city if t.upper() not in {prov, postal})
    return street, city, prov, postal


def ensure_need_appearances(writer: PdfWriter) -> None:
    try:
        writer._root_object["/AcroForm"].update(
            {NameObject("/NeedAppearances"): BooleanObject(True)}
        )
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Extraction routines (ported from Azure Function)

def extract_mrp_data(text: str) -> Dict[str, str]:
    result: Dict[str, str] = {}
    section = re.search(
        r"New Payments.*?(.*?)Total of Payment Activity",
        text,
        re.DOTALL | re.IGNORECASE,
    )
    if section:
        dates = re.findall(r"([A-Z][a-z]{2}\s?\d{1,2})", section.group(1))
        if len(dates) >= 2:
            transaction = re.sub(r"([A-Za-z]{3})(\d{1,2})", r"\1 \2", dates[-2].strip())
            posting = re.sub(r"([A-Za-z]{3})(\d{1,2})", r"\1 \2", dates[-1].strip())
            result["last_payment_date"] = (
                f"Payment Transaction Date: {transaction}, Posting Date: {posting}"
            )
    date_range = re.search(r"(\w{3}\s?\d{1,2},\s?\d{4})\s+(\w{3}\s?\d{1,2},\s?\d{4})", text)
    if date_range:
        result["opening_date"], result["closing_date"] = date_range.groups()
    return result


def extract_mrc_data(text: str) -> Dict[str, str]:
    result: Dict[str, str] = {}
    section = re.search(
        r"New Transactions for.*?(.*?)Total of New Transactions",
        text,
        re.DOTALL | re.IGNORECASE,
    )
    if section:
        dates = re.findall(r"([A-Z][a-z]{2}\s?\d{1,2})", section.group(1))
        if len(dates) >= 2:
            transaction = re.sub(r"([A-Za-z]{3})(\d{1,2})", r"\1 \2", dates[-2].strip())
            posting = re.sub(r"([A-Za-z]{3})(\d{1,2})", r"\1 \2", dates[-1].strip())
            result["last_charge_date"] = (
                f"Transaction Date: {transaction}, Posting Date: {posting}"
            )
    date_range = re.search(r"(\w{3}\s?\d{1,2},\s?\d{4})\s+(\w{3}\s?\d{1,2},\s?\d{4})", text)
    if date_range:
        result["opening_date"], result["closing_date"] = date_range.groups()
    return result


def extract_cbr_data(text: str) -> Dict[str, object]:
    result: Dict[str, object] = {}
    name_match = re.search(r"This completes the file for (.+)", text)
    if name_match:
        result["name"] = name_match.group(1).strip()
    phone_match = re.search(r"\b\d{10}\b", text)
    if phone_match:
        result["phone"] = phone_match.group(0)

    names_in_order: List[Dict[str, str]] = []
    subject_match = re.search(r"Subject\s+([A-Z]+)\s+([A-Z/]+(?: [A-Z/]+)*)\s+\d", text)
    if subject_match:
        surname = subject_match.group(1).strip()
        given = subject_match.group(2).strip()
        names_in_order.append(
            {"surname": surname, "given_name": given, "full_name": f"{given} {surname}"}
        )
    else:
        alt = re.search(
            r"Surname\s+([A-Z]+)\s+Given Name\(s\)\s+([A-Z/ ]+?)\s+Soc\.Ins\.No",
            text,
        )
        if alt:
            surname = alt.group(1).strip()
            given = alt.group(2).strip()
            names_in_order.append(
                {"surname": surname, "given_name": given, "full_name": f"{given} {surname}"}
            )

    aka_block = re.search(
        r"X-Ref AKA(.+?)(?=On File|Last|Current|Street|Birth|Given|Surname|Reference|$)",
        text,
    )
    if aka_block:
        for entry in re.split(r"\bAKA\b", aka_block.group(1)):
            parts = entry.strip().split()
            if len(parts) >= 2:
                surname = parts[0]
                given = " ".join(parts[1:])
                names_in_order.append(
                    {"surname": surname, "given_name": given, "full_name": f"{given} {surname}"}
                )

    if names_in_order:
        result["name"] = names_in_order[0]["full_name"]
        result["akas"] = names_in_order[1:] if len(names_in_order) > 1 else []

    return result


def extract_mrs_data(text: str) -> Dict[str, object]:
    result: Dict[str, object] = {}
    lines = [normalize_unicode(line) for line in text.splitlines()]
    for idx in range(len(lines) - 3, -1, -1):
        possible_name = lines[idx].strip()
        addr1 = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
        addr2 = lines[idx + 2].strip() if idx + 2 < len(lines) else ""
        if re.match(r"^(APT|UNIT|SUITE)\b", possible_name.upper()):
            if idx - 1 >= 0:
                actual_name = lines[idx - 1].strip()
                address = f"{possible_name}, {addr1}, {addr2}"
            else:
                continue
        else:
            actual_name = possible_name
            address = f"{addr1}, {addr2}"
        postal_candidate = re.search(r"[A-Z]\d[A-Z]\s?\d[A-Z]\d", addr2.upper())
        if (
            len(actual_name.split()) >= 2
            and re.search(r"\d{3,}", address)
            and postal_candidate
        ):
            result["name"] = actual_name
            result["address"] = address
            break

    result["account_type"] = "Charge Card" if "charge card" in text.lower() else "Credit Card"

    for line in lines:
        if "Statement of Account" in line:
            break
        if re.search(r"\bcard\b", line, re.IGNORECASE) and not re.search(
            r"\d{3}[-\s]?\d{3}[-\s]?\d{4}", line
        ):
            result["card_name"] = line.strip()
            break

    def _extract_amount(pattern: str) -> Optional[float]:
        match = re.search(pattern, text)
        if match:
            return float(match.group(1).replace(",", ""))
        return None

    result["previous_balance"] = _extract_amount(r"Previous Balance\s*\$([\d,]+\.\d{2})")
    result["interest"] = _extract_amount(r"Plus Interest\s*\$([\d,]+\.\d{2})")
    result["fees"] = _extract_amount(r"Plus Fees\s*\$([\d,]+\.\d{2})")
    result["payments"] = _extract_amount(r"Less Payments\s*\$([\d,]+\.\d{2})")
    total = _extract_amount(r"New Balance\s*\$([\d,]+\.\d{2})")
    if total is not None:
        result["total_indebtedness"] = total
        result["outstanding_balance"] = f"{total:,.2f}"

    account_match = re.search(r"XXXX XXXXX\d (\d{5})", text)
    if account_match:
        last_five = account_match.group(1)
        result["subject_line"] = f"XX1 {last_five}"
        pattern = rf"{last_five}\s+(\w{{3}}\.?\s?\d{{1,2}},\s?\d{{4}})\s+(\w{{3}}\.?\s?\d{{1,2}},\s?\d{{4}})"
        dates = re.search(pattern, text)
        if dates:
            result["opening_date"], result["closing_date"] = dates.groups()

    interest_match = re.search(r"Annual Interest Rate\s*[:\-]?\s*(\d+\.\d+)%", text)
    if not interest_match:
        interest_match = re.search(r"Purchases\s+[\d\.]+%\s+[\d\.]+\s+(\d+\.\d+)%", text)
    if interest_match:
        result["annual_interest_rate"] = interest_match.group(1)

    return result


def extract_basic_info(text: str) -> Dict[str, str]:
    result: Dict[str, str] = {}
    name_match = re.search(r"This completes the file for (.+)", text)
    if name_match:
        result["name"] = name_match.group(1).strip()
    phone_match = re.search(r"\b\d{10}\b", text)
    if phone_match:
        result["phone"] = phone_match.group(0)
    return result


FILE_EXTRACTORS = {
    "MRP": extract_mrp_data,
    "MRC": extract_mrc_data,
    "MRS": extract_mrs_data,
    "CBR": extract_cbr_data,
}


def extract_bc_claims_data(pdf_path: Path) -> Dict[str, object]:
    file_type = None
    upper_name = pdf_path.name.upper()
    for key in FILE_EXTRACTORS:
        if key in upper_name:
            file_type = key
            break

    aggregated: Dict[str, object] = {
        "akas": [],
    }

    with pdfplumber.open(str(pdf_path)) as pdf:
        text = "\n".join((page.extract_text() or "") for page in pdf.pages)

    if file_type and file_type in FILE_EXTRACTORS:
        aggregated.update(FILE_EXTRACTORS[file_type](text))

    aggregated.update(extract_basic_info(text))
    aggregated["name1"] = aggregated.get("name", "UNKNOWN")
    aggregated["address1"] = aggregated.get("address", "UNKNOWN")
    aggregated.setdefault("file_type", file_type or "UNKNOWN")
    return aggregated


def merge_claim_data(base: Dict[str, object], new: Dict[str, object]) -> Dict[str, object]:
    def is_meaningful(value) -> bool:
        if value in (None, ""):
            return False
        if isinstance(value, str) and value.upper() in {"UNKNOWN", "UNKNOWN NAME", "UNKNOWN ADDRESS"}:
            return False
        if isinstance(value, list):
            return bool(value)
        return True

    for key, value in new.items():
        if not is_meaningful(value):
            continue
        if isinstance(value, list):
            base[key] = value
        elif isinstance(value, dict):
            base[key] = merge_claim_data(base.get(key, {}), value)
        else:
            base[key] = value
    return base


# ---------------------------------------------------------------------------
# Formatting for templates

def format_data_for_template(data: Dict[str, object]) -> Dict[str, str]:
    formatted_name = data.get("name", "UNKNOWN NAME")
    akas = data.get("akas") or []
    if akas:
        aka_part = " aka ".join(f"{aka['given_name']} {aka['surname']}" for aka in akas)
        formatted_name = f"{formatted_name} aka {aka_part}"

    def format_date(date_str: Optional[str]) -> str:
        if not date_str:
            return "UNKNOWN"
        return re.sub(r"([A-Za-z]{3})(\d{2})", r"\1 \2", str(date_str))

    def format_with_full_month(date_str: str, year: int) -> str:
        try:
            dt = datetime.strptime(f"{date_str} {year}", "%b %d %Y")
            return dt.strftime("%B %d, %Y")
        except ValueError:
            return date_str

    def get_year_from_closing(closing_date: Optional[str]) -> Optional[int]:
        if not closing_date:
            return None
        try:
            return datetime.strptime(closing_date.replace(" ", ""), "%b%d,%Y").year
        except ValueError:
            try:
                return datetime.strptime(closing_date, "%b %d, %Y").year
            except ValueError:
                return None

    closing_date = data.get("closing_date")
    year = get_year_from_closing(closing_date)

    raw_payment = None
    if isinstance(data.get("last_payment_date"), str):
        raw_payment = data["last_payment_date"].replace("Payment Transaction Date:", "").split(",")[0].strip()
    last_payment = (
        format_with_full_month(raw_payment, year) if raw_payment and year else (raw_payment or "UNKNOWN")
    )

    raw_charge = None
    if isinstance(data.get("last_charge_date"), str):
        raw_charge = data["last_charge_date"].replace("Transaction Date:", "").split(",")[0].strip()
    last_charge = (
        format_with_full_month(raw_charge, year) if raw_charge and year else (raw_charge or "UNKNOWN")
    )

    opening = format_date(data.get("opening_date"))
    closing = format_date(closing_date)

    formatted = {
        "SIMPLE NAME INSERT": formatted_name,
        "ADRESS INSERT": data.get("address", "UNKNOWN ADDRESS"),
        "CREDIT CARD / CHARGE CARD / PERSONAL LOAN INSERT": data.get("account_type", "UNKNOWN ACCOUNT TYPE"),
        "CARD TYPE INSERT": data.get("card_name", "UNKNOWN CARD"),
        "DATE OF LAST CHARGE INSERT": last_charge or "UNKNOWN",
        "DATE OF LAST PAYMENT INSERT": last_payment or "UNKNOWN",
        "INTEREST RATE INSERT": data.get("annual_interest_rate", "UNKNOWN"),
        "DATE RANGE INSERT": f"{opening} - {closing}" if opening != "UNKNOWN" and closing != "UNKNOWN" else "UNKNOWN RANGE",
        "SMALLDINSERT": to_money_str(data.get("previous_balance")),
        "ACCRUED INTEREST INSERT": to_money_str(data.get("interest")),
        "FINSERT": to_money_str(data.get("fees") or 0),
        "PAYMENTS INSERT": to_money_str(data.get("payments") or 0),
        "DEBT INSERT": to_money_str(data.get("total_indebtedness") or data.get("outstanding_balance")),
        "DEMAND LETTER DATE INSERT": datetime.now().strftime("%b %d, %Y"),
        "TELINSERT": data.get("phone", ""),
        "CARD INSERT": data.get("subject_line", ""),
        "AMOUNT INSERT": to_money_str(data.get("outstanding_balance")),
    }
    return formatted


# ---------------------------------------------------------------------------
# Schedule A document generation

def replace_placeholder_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return
    new_text = full_text.replace(placeholder, replacement)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_text_everywhere(doc: Document, placeholder: str, replacement: str) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, placeholder, replacement)


def create_schedule_a_document(template_path: Path, data: Dict[str, str]) -> bytes:
    doc = Document(str(template_path))
    for placeholder, value in data.items():
        replace_text_everywhere(doc, placeholder, str(value))
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Notice of Claim PDF generation

def build_pdf_field_values_from_data(data: Dict[str, str]) -> Dict[str, str]:
    simple_name = data.get("SIMPLE NAME INSERT", "").strip()
    address_full = data.get("ADRESS INSERT", data.get("ADDRESS INSERT", "")).strip()
    date_range = data.get("DATE RANGE INSERT", "").strip()
    rate = data.get("INTEREST RATE INSERT", "").strip()
    rate_display = rate if rate.endswith("%") or not rate else f"{rate}%"
    debt_principal = to_money_str(data.get("SMALLDINSERT"))
    accrued_interest = to_money_str(data.get("ACCRUED INTEREST INSERT"))
    payments = to_money_str(data.get("PAYMENTS INSERT"))
    total_debt = to_money_str(data.get("DEBT INSERT"))

    street, city, prov, postal = parse_bc_address(address_full)
    phone = normalize_phone(data.get("TELINSERT", ""))

    part1 = f"Principle debt owing on last statement from {date_range}".strip()
    account_type = data.get("CREDIT CARD / CHARGE CARD / PERSONAL LOAN INSERT", "").strip()
    card_type = data.get("CARD TYPE INSERT", "").strip()
    if account_type or card_type:
        part1 += f" ({account_type}: {card_type})"
    part2 = (
        f"Contractual interest at the rate of {rate_display or 'XX.XX%'} per annum or as may be as amended "
        f"from time to time from {date_range}"
    )
    part3 = "Payments/Credits applied"

    try:
        base_total = Decimal(str(total_debt).replace("$", "").replace(",", ""))
        total2 = f"${(base_total + Decimal(236)):,.2f}"
    except InvalidOperation:
        total2 = total_debt

    return {
        "defendant": simple_name,
        "addressd": street,
        "cityd": city,
        "provd": prov,
        "zipd": postal,
        "phoned": phone,
        "where": city,
        "provo": prov,
        "when": date_range,
        "part1": part1,
        "amounta": debt_principal,
        "part2": part2,
        "amountb": accrued_interest,
        "part3": part3,
        "amountc": payments,
        "total": total_debt,
        "total2": total2,
    }


def create_notice_of_claim_pdf(template_path: Path, data: Dict[str, str]) -> bytes:
    reader = PdfReader(str(template_path))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    field_values = build_pdf_field_values_from_data(data)
    for page in writer.pages:
        writer.update_page_form_field_values(page, field_values)
    ensure_need_appearances(writer)
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Core workflow

@dataclass
class ClaimsInputs:
    mrp: Path
    mrc: Path
    mrs: Path
    cbr: Path


def resolve_inputs(args: argparse.Namespace) -> ClaimsInputs:
    def resolve_path(cli_value: Optional[str], patterns: Iterable[str]) -> Path:
        if cli_value:
            path = Path(cli_value).expanduser().resolve()
            if not path.exists():
                raise FileNotFoundError(f"File not found: {path}")
            return path
        path = auto_pick_file(INPUT_DIR, patterns)
        if not path:
            joined = " or ".join(patterns)
            raise FileNotFoundError(
                f"Missing required file matching pattern {joined!r} in {INPUT_DIR}"
            )
        return path.resolve()

    return ClaimsInputs(
        mrp=resolve_path(args.mrp, FILE_TYPE_PATTERNS["MRP"]),
        mrc=resolve_path(args.mrc, FILE_TYPE_PATTERNS["MRC"]),
        mrs=resolve_path(args.mrs, FILE_TYPE_PATTERNS["MRS"]),
        cbr=resolve_path(args.cbr, FILE_TYPE_PATTERNS["CBR"]),
    )


def aggregate_claims_data(paths: ClaimsInputs) -> Dict[str, object]:
    combined: Dict[str, object] = {}
    for pdf_path in [paths.mrp, paths.mrc, paths.mrs, paths.cbr]:
        extracted = extract_bc_claims_data(pdf_path)
        merge_claim_data(combined, extracted)
    return combined


def save_bytes(data: bytes, destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with open(destination, "wb") as handle:
        handle.write(data)
    return destination


def run_workflow(args: argparse.Namespace) -> Dict[str, str]:
    inputs = resolve_inputs(args)
    combined = aggregate_claims_data(inputs)
    formatted = format_data_for_template(combined)

    schedule_bytes = create_schedule_a_document(SCHEDULE_TEMPLATE, formatted)
    notice_bytes = create_notice_of_claim_pdf(NOTICE_TEMPLATE, formatted)

    defendant = formatted.get("SIMPLE NAME INSERT", "Defendant")
    output_name = safe_filename(defendant)

    schedule_path = OUTPUT_DIR / f"{output_name} - Schedule A.docx"
    notice_path = OUTPUT_DIR / f"{output_name} - Notice of Claim.pdf"

    save_bytes(schedule_bytes, schedule_path)
    save_bytes(notice_bytes, notice_path)

    return {
        "defendant": defendant,
        "address": formatted.get("ADRESS INSERT", ""),
        "card_reference": formatted.get("CARD INSERT", ""),
        "amount": formatted.get("AMOUNT INSERT", ""),
        "schedule_path": str(schedule_path),
        "notice_path": str(notice_path),
    }


# ---------------------------------------------------------------------------
# CLI

def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate BC Claims Schedule A and Notice of Claim from source PDFs."
    )
    parser.add_argument("--mrp", help="Path to the MRP (payments) PDF.")
    parser.add_argument("--mrc", help="Path to the MRC (charges) PDF.")
    parser.add_argument("--mrs", help="Path to the MRS (statement) PDF.")
    parser.add_argument("--cbr", help="Path to the CBR (credit bureau report) PDF.")
    parser.add_argument("--output-dir", help="Override output directory.")
    return parser.parse_args(argv)


def main(argv: Optional[List[str]] = None) -> None:
    args = parse_args(argv)
    global OUTPUT_DIR
    if args.output_dir:
        OUTPUT_DIR = Path(args.output_dir).expanduser().resolve()
    summary = run_workflow(args)
    print("Generated BC Claims documents:")
    print(f"  Defendant        : {summary['defendant']}")
    print(f"  Address          : {summary['address']}")
    print(f"  Card Reference   : {summary['card_reference']}")
    print(f"  Outstanding Amt  : {summary['amount']}")
    print(f"  Schedule A       : {summary['schedule_path']}")
    print(f"  Notice of Claim  : {summary['notice_path']}")


if __name__ == "__main__":
    main()

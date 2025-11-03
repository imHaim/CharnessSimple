# -*- coding: utf-8 -*-
"""Generate BC demand letter documents from local assets."""

from __future__ import annotations

import argparse
import io
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional, Tuple

import pdfplumber
from docx import Document

MODULE_DIR = Path(__file__).resolve().parent
ASSET_DIR = MODULE_DIR.parent / "assets" / "demand_letter"
DEFAULT_TEMPLATE = ASSET_DIR / "Demand-Letter-BC.docx"


ACCOUNT_PATTERN = re.compile(r"XXXX XXXXX\d (\d{5})")
BALANCE_PATTERN = re.compile(r"(?:New Balance|Balance)\s+\$(\d{1,3}(?:,\d{3})*(?:\.\d{2}))")


def normalize_unicode(value: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\ufffd": "",
    }
    for raw, replacement in replacements.items():
        value = value.replace(raw, replacement)
    return value


def safe_filename(value: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", value)
    sanitized = re.sub(r"_{2,}", "_", sanitized).strip(" .")
    return sanitized or "output"


def extract_name_and_address_from_bottom(pdf_bytes: bytes) -> Tuple[str, str]:
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text() or ""

    lines = [normalize_unicode(line.strip()) for line in text.splitlines()]

    for idx in range(len(lines) - 3, -1, -1):
        name_line = lines[idx].strip()
        addr1 = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
        addr2 = lines[idx + 2].strip() if idx + 2 < len(lines) else ""
        has_name = len(name_line.split()) >= 2
        has_street = bool(re.search(r"\d", addr1))
        has_postal = bool(re.search(r"[A-Z]\d[A-Z]\s?\d[A-Z]\d", addr2.upper()))
        if has_name and has_street and has_postal:
            address = ", ".join(part for part in (addr1, addr2) if part)
            return name_line, address

    for idx in range(max(0, len(lines) - 5), len(lines)):
        line = lines[idx]
        if re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+", line) or re.match(r"^Mr\s+[A-Z][a-z]+", line, re.IGNORECASE):
            return line, "Address not found"

    return "", ""


def extract_data_from_pdf(pdf_path: Path) -> dict[str, str]:
    pdf_bytes = pdf_path.read_bytes()
    name, address = extract_name_and_address_from_bottom(pdf_bytes)

    subject_line = ""
    outstanding_balance = ""

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        page = pdf.pages[0]
        text = normalize_unicode(page.extract_text() or "")

    account_match = ACCOUNT_PATTERN.search(text)
    if account_match:
        subject_line = f"XX1 {account_match.group(1)}"

    balance_match = BALANCE_PATTERN.search(text)
    if balance_match:
        outstanding_balance = balance_match.group(1)

    today = datetime.now().strftime("%B %d, %Y")

    return {
        "CARD INSERT": subject_line or "XX1 XXXXX",
        "AMOUNT INSERT": outstanding_balance or "0.00",
        "NAME INSERT": name or "NAME NOT FOUND",
        "ADDRESS INSERT": address or "ADDRESS NOT FOUND",
        "DATE INSERT": today,
    }


def replace_text(paragraphs, replacements: dict[str, str]) -> None:
    for paragraph in paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)


def create_demand_letter_doc(template_path: Path, data: dict[str, str]) -> bytes:
    if template_path.exists():
        doc = Document(template_path)
    else:
        doc = Document()
        doc.add_paragraph("WITHOUT PREJUDICE")
        doc.add_paragraph(data["DATE INSERT"])
        doc.add_paragraph(data["NAME INSERT"])
        doc.add_paragraph(data["ADDRESS INSERT"])
        doc.add_paragraph("")
        doc.add_paragraph(
            f"RE: Amex Bank of Canada Account No. Ending with < XXXX XXXX {data['CARD INSERT']}>"
        )
        doc.add_paragraph(
            f"        Outstanding Balance of Latest Statement: ${data['AMOUNT INSERT']}"
        )
        doc.add_paragraph("__________________________________________________________________________")
        doc.add_paragraph("")
        doc.add_paragraph(f"Dear {data['NAME INSERT']},")
        doc.add_paragraph("")
        doc.add_paragraph("We have been retained to collect the above-noted claim.")
        doc.add_paragraph("")
        doc.add_paragraph(
            "On behalf of our client, we hereby make formal and final demand for immediate payment in full. "
            "Your remittance should be made payable to Charness, Charness & Charness, In Trust, and mailed or delivered to us."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "Unless full payment is received or satisfactory arrangements for payment have been made within 10 days from "
            "the date of this letter, legal proceedings may be initiated against you without further notice."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Yours very truly,")
        doc.add_paragraph("")
        doc.add_paragraph("CHARNESS, CHARNESS & CHARNESS LLP")
        doc.add_paragraph("")
        doc.add_paragraph("Per: _________________________________")
        doc.add_paragraph("Jordan W. Charness")
        doc.add_paragraph("JWC/initials")

    replace_text(doc.paragraphs, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell.paragraphs, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def fill_demand_letter(pdf_path: Path, template_path: Optional[Path] = None, output_dir: Optional[Path] = None):
    template = template_path or DEFAULT_TEMPLATE
    out_dir = output_dir or (MODULE_DIR.parent / "output")
    out_dir.mkdir(parents=True, exist_ok=True)

    data = extract_data_from_pdf(pdf_path)
    letter_bytes = create_demand_letter_doc(template, data)

    output_path = out_dir / f"{safe_filename(data['NAME INSERT'])} - BC Demand Letter.docx"
    output_path.write_bytes(letter_bytes)
    return output_path, data


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a BC demand letter from a source PDF.")
    parser.add_argument("pdf", help="Path to the MRS/statement PDF.")
    parser.add_argument("--template", help="Override template DOCX path.")
    parser.add_argument("--output-dir", help="Override output directory.")
    return parser.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> None:  # pragma: no cover
    args = parse_args(argv)
    pdf_path = Path(args.pdf).expanduser().resolve()
    template = Path(args.template).expanduser().resolve() if args.template else DEFAULT_TEMPLATE
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else MODULE_DIR.parent / "output"
    out, summary = fill_demand_letter(pdf_path, template, output_dir)
    print(f"Saved demand letter to: {out}")
    print(summary)


if __name__ == "__main__":  # pragma: no cover
    main()


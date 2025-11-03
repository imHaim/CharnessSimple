# -*- coding: utf-8 -*-
"""Automate the BC Small Claims Application for Default Order form.

This script extracts key values from a Schedule A (DOCX) and a Notice of Claim
PDF, then populates the Application for Default Order PDF template. Paths are
resolved relative to this project so the script can run on any machine without
manual edits.
"""

from __future__ import annotations

import argparse
import re
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject

MODULE_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT_DIR = MODULE_DIR / "input"
DEFAULT_TEMPLATE_NAME = "Application for Default Order Template.pdf"
ASSET_TEMPLATE = MODULE_DIR.parent / "assets" / "default_judgment" / DEFAULT_TEMPLATE_NAME
DEFAULT_OUTPUT_DIR = MODULE_DIR.parent / "output"

DEFAULT_EXPENSES = Decimal("236.00")  # $156 filing + $80 service
DEFAULT_APPLICATION_FEE = Decimal("25.00")

MONEY_RE = re.compile(r"\$(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)")
TOTAL_RE = re.compile(r"TOTAL\s*\$\s*([\d,]+\.\d{2})", re.I)

FIELD_KEYS = {
    "defendant_1": "defendant",
    "defendant_2": "def",
    "left_text": "terms",
    "amount_a": "claim",
    "amount_b": "expenses",
    "amount_c": "interest",
    "amount_e": "fee",
    "registry_loc": "location",
}


# ---------------------------------------------------------------------------
# Parsing helpers

def parse_claim_schedule_a(docx_path: Path) -> dict[str, str]:
    """Extract defendant, interest rate, date range, principal, and total claimed."""
    doc = Document(docx_path)
    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]

    table_rows: list[list[str]] = []
    table_strings: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells)
            if any(cells):
                table_strings.append(" ".join(filter(None, cells)))

    text = "\n".join(paragraph_texts + table_strings)

    defendant_match = re.search(
        r"([A-Z][A-Z/ '\-]+?)(?=\s*\(the[^)]*Defendant[^)]*\))",
        text,
        flags=re.IGNORECASE,
    )
    defendant_full = (
        normalize_unicode(defendant_match.group(1)).strip()
        if defendant_match
        else ""
    )

    interest_match = re.search(
        r"(\d{1,2}\.\d{2})\s*%\s*per\s*annum", text, flags=re.IGNORECASE
    )
    interest_rate = interest_match.group(1) if interest_match else ""

    principal = ""
    date_range = ""
    for cells in table_rows:
        if not cells:
            continue
        heading = cells[0].strip()
        if heading.lower().startswith("debt owing on the last statement"):
            date_match = re.search(r"\(([^)]+)\)", heading)
            if date_match:
                date_range = normalize_unicode(date_match.group(1)).strip()
            if len(cells) > 1:
                money_match = MONEY_RE.search(cells[1])
                if money_match:
                    principal = f"${money_match.group(1)}"
            break

    total_match = TOTAL_RE.search(text)
    total_claimed = f"${total_match.group(1)}" if total_match else ""
    if not total_claimed:
        fallback = re.search(
            r"Total\s+Indebtedness.*?\$([\d,]+\.\d{2})",
            text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if fallback:
            total_claimed = f"${fallback.group(1)}"

    return {
        "defendant_full": defendant_full,
        "interest_rate": interest_rate,
        "date_range": date_range,
        "principal_last_stmt": principal,
        "total_claimed": total_claimed,
    }


def read_notice_interest_and_registry(pdf_path: Path) -> dict[str, str]:
    """Read registry location and interest amount from the Notice of Claim PDF."""
    reader = PdfReader(str(pdf_path))
    fields = reader.get_fields() or {}

    def get_like(*keys: str) -> str:
        for key in keys:
            if key in fields and fields[key].get("/V"):
                return str(fields[key]["/V"]).strip()
        lowered = [key.lower() for key in keys]
        for key, value in fields.items():
            raw_val = value.get("/V")
            if not raw_val:
                continue
            haystack = (
                (value.get("/T") or "")
                + " "
                + (value.get("/TU") or "")
                + " "
                + (key or "")
            ).lower()
            if any(fragment in haystack for fragment in lowered):
                return str(raw_val).strip()
        return ""

    return {
        "registry_location": get_like(
            "location", "registry location", "court location"
        ),
        "interest_amount": get_like(
            "amountb", "interest amount", "amount c", "interest"
        ),
    }


# ---------------------------------------------------------------------------
# Text helpers

def normalize_unicode(value: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\ufffd": "-",
    }
    for raw, replacement in replacements.items():
        value = value.replace(raw, replacement)
    return value


def build_left_text_three_blocks(rate: str, date_range: str) -> str:
    """Construct the left narrative box as three CRLF-separated paragraphs."""
    rate_text = (rate or "XX.XX").strip()
    date_text = (date_range or "Month Day, Year to Month Day, Year").strip()
    return (
        "$156 FILING FEE AND $80 PERSONAL SERVICE"
        + "\r\n\r\n"
        + f"Contractual interest at the rate of {rate_text}% per annum from {date_text}"
        + "\r\n\r\n"
        + "$25 APPLICATION FEE"
    )


def extract_primary_name(defendant_full: str) -> tuple[str, str]:
    base = (defendant_full or "").split(" aka ")[0].strip()
    parts = [part.strip(",") for part in re.split(r"\s+", base) if part]
    if len(parts) >= 2:
        return parts[0], parts[-1]
    if parts:
        return parts[0], ""
    return "UNKNOWN", ""


def safe_filename(text: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", text)
    sanitized = re.sub(r"_{2,}", "_", sanitized).strip(" .")
    return sanitized or "output"


def make_outfile(defendant_full: str, output_dir: Path) -> Path:
    first, last = extract_primary_name(defendant_full)
    if last:
        filename = f"{last}, {first} - Filed Default Order.pdf"
    else:
        filename = f"{first} - Filed Default Order.pdf"
    return output_dir / safe_filename(filename)


# ---------------------------------------------------------------------------
# PDF helpers

def ensure_need_appearances(writer: PdfWriter) -> None:
    try:
        writer._root_object["/AcroForm"].update(
            {NameObject("/NeedAppearances"): BooleanObject(True)}
        )
    except KeyError:
        pass


def write_field_all_pages(pdf_writer: PdfWriter, field_key: str, value: str) -> None:
    if not field_key or value is None:
        return
    for page in pdf_writer.pages:
        if "/Annots" in page:
            pdf_writer.update_page_form_field_values(page, {field_key: value})


def put_amount(pdf_writer: PdfWriter, field_key: str, value: str) -> None:
    """Populate a numeric field with currency and digits-only fallbacks."""
    if not field_key or not value:
        return
    write_field_all_pages(pdf_writer, field_key, value)
    digits_only = re.sub(r"[^\d.]", "", value)
    if digits_only and digits_only != value:
        write_field_all_pages(pdf_writer, field_key, digits_only)


# ---------------------------------------------------------------------------
# Numeric helpers

def to_decimal(raw_value: str | Decimal | None) -> Optional[Decimal]:
    if raw_value in (None, ""):
        return None
    stripped = re.sub(r"[^\d.\-]", "", str(raw_value))
    try:
        return Decimal(stripped)
    except InvalidOperation:
        return None


def format_currency(amount: Decimal) -> str:
    return f"${amount:,.2f}"


def compute_principal_amount(
    parsed_values: dict[str, str], interest_amount: str
) -> str:
    principal = to_decimal(parsed_values.get("principal_last_stmt"))
    if principal is not None:
        return format_currency(principal)

    total_claimed = to_decimal(parsed_values.get("total_claimed"))
    if total_claimed is None:
        return ""

    interest = to_decimal(interest_amount) or Decimal("0.00")
    principal = total_claimed - DEFAULT_EXPENSES - DEFAULT_APPLICATION_FEE - interest
    if principal <= 0:
        return ""
    return format_currency(principal)


# ---------------------------------------------------------------------------
# Core workflow

def fill_default_order(
    claim_docx: Path,
    notice_pdf: Path,
    template_pdf: Path,
    output_dir: Path,
    registry_location_override: str = "",
) -> tuple[Path, dict[str, str]]:
    values = parse_claim_schedule_a(claim_docx)
    notice = read_notice_interest_and_registry(notice_pdf)

    left_block = build_left_text_three_blocks(
        values.get("interest_rate", ""), values.get("date_range", "")
    )

    amount_a = compute_principal_amount(values, notice.get("interest_amount", ""))
    amount_b = format_currency(DEFAULT_EXPENSES)
    amount_c = notice.get("interest_amount") or ""
    amount_e = format_currency(DEFAULT_APPLICATION_FEE)

    reader = PdfReader(str(template_pdf))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)

    defendant_full = values.get("defendant_full", "")
    if defendant_full:
        write_field_all_pages(writer, FIELD_KEYS["defendant_1"], defendant_full)
        write_field_all_pages(writer, FIELD_KEYS["defendant_2"], defendant_full)

    write_field_all_pages(writer, FIELD_KEYS["left_text"], left_block)
    put_amount(writer, FIELD_KEYS["amount_a"], amount_a)
    put_amount(writer, FIELD_KEYS["amount_b"], amount_b)
    put_amount(writer, FIELD_KEYS["amount_c"], amount_c)
    put_amount(writer, FIELD_KEYS["amount_e"], amount_e)

    registry_location = notice.get("registry_location") or registry_location_override
    if registry_location:
        write_field_all_pages(writer, FIELD_KEYS["registry_loc"], registry_location)

    ensure_need_appearances(writer)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = make_outfile(defendant_full, output_dir)
    with open(output_path, "wb") as file_obj:
        writer.write(file_obj)

    summary = {
        "defendant": defendant_full or "(not found)",
        "interest_rate": values.get("interest_rate", ""),
        "interest_period": values.get("date_range", ""),
        "amount_a": amount_a or "(blank)",
        "amount_b": amount_b,
        "amount_c": amount_c or "(blank)",
        "amount_e": amount_e,
        "registry_location": registry_location or "(blank)",
    }
    return output_path, summary


# ---------------------------------------------------------------------------
# CLI helpers

def auto_pick_file(directory: Path, patterns: Iterable[str]) -> Optional[Path]:
    if not directory.exists():
        return None
    for pattern in patterns:
        matches = sorted(
            directory.glob(pattern),
            key=lambda candidate: candidate.stat().st_mtime,
            reverse=True,
        )
        if matches:
            return matches[0]
    return None


def resolve_inputs_from_args(args: argparse.Namespace) -> tuple[Path, Path, Path, Path]:
    claim_docx = Path(args.claim_docx).expanduser().resolve() if args.claim_docx else None
    notice_pdf = Path(args.notice_pdf).expanduser().resolve() if args.notice_pdf else None
    template_pdf = (
        Path(args.template_pdf).expanduser().resolve() if args.template_pdf else None
    )
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else None

    if claim_docx is None:
        claim_docx = auto_pick_file(
            DEFAULT_INPUT_DIR, ("*Schedule*.docx", "*.docx")
        )
    if notice_pdf is None:
        notice_pdf = auto_pick_file(
            DEFAULT_INPUT_DIR, ("*Notice*.pdf", "*.pdf")
        )
    if template_pdf is None:
        template_pdf = ASSET_TEMPLATE
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR

    missing = [
        ("Schedule A DOCX", claim_docx),
        ("Notice of Claim PDF", notice_pdf),
        ("Default Order template PDF", template_pdf),
    ]
    for label, path in missing:
        if path is None or not path.exists():
            raise FileNotFoundError(f"Unable to locate {label}. Checked path: {path}")

    return claim_docx, notice_pdf, template_pdf, output_dir


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Populate the BC Small Claims Application for Default Order form."
    )
    parser.add_argument("--claim-docx", help="Path to the filled Schedule A DOCX file.")
    parser.add_argument("--notice-pdf", help="Path to the filed Notice of Claim PDF.")
    parser.add_argument(
        "--template-pdf",
        help="Path to the Application for Default Order PDF template.",
    )
    parser.add_argument(
        "--output-dir",
        help="Directory to store the populated PDF (default: ./output).",
    )
    parser.add_argument(
        "--registry-location",
        default="",
        help="Override registry location if the Notice is missing the field.",
    )
    return parser.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> None:
    args = parse_args(argv)
    claim_docx, notice_pdf, template_pdf, output_dir = resolve_inputs_from_args(args)

    output_path, summary = fill_default_order(
        claim_docx=claim_docx,
        notice_pdf=notice_pdf,
        template_pdf=template_pdf,
        output_dir=output_dir,
        registry_location_override=args.registry_location,
    )

    print(f"Saved default order to: {output_path}")
    print("Summary:")
    for key, value in summary.items():
        print(f"  {key.replace('_', ' ').title()}: {value}")


if __name__ == "__main__":
    main()

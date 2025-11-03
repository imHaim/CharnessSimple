# -*- coding: utf-8 -*-
"""Populate the BC Small Claims Notice of Withdrawal (dismissal) form.

The script reads the defendant name from a Schedule A DOCX and (optionally) the
registry file number from a Notice of Claim PDF, then fills the PDF template.
Relative paths are used so the workflow runs on any machine without edits.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject

MODULE_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT_DIR = MODULE_DIR / "input"
DEFAULT_TEMPLATE_NAME = "Notice of Withdrawal template.pdf"
ASSET_TEMPLATE = MODULE_DIR.parent / "assets" / "dismissal" / DEFAULT_TEMPLATE_NAME
DEFAULT_OUTPUT_DIR = MODULE_DIR.parent / "output"

DEFENDANT_REGEX = r"([A-Z][A-Z/ '\-]+?)(?=\s*\(the[^)]*Defendant[^)]*\))"


# ---------------------------------------------------------------------------
# Text helpers

def normalize_unicode(value: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\ufffd": "-",
    }
    for raw, replacement in replacements.items():
        value = value.replace(raw, replacement)
    return value


def safe_filename(name: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", name)
    sanitized = re.sub(r"_{2,}", "_", sanitized).strip(" .")
    return sanitized or "output"


# ---------------------------------------------------------------------------
# Parsing helpers

def extract_defendant_from_claim(docx_path: Path) -> str:
    try:
        document = Document(docx_path)
    except Exception as exc:  # pragma: no cover - defensive
        print(f"Could not open DOCX: {exc}")
        return "UNKNOWN DEFENDANT"

    paragraph_texts = [p.text for p in document.paragraphs if p.text]
    table_strings = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_strings.append(" ".join(filter(None, cells)))

    combined = "\n".join(paragraph_texts + table_strings)
    match = re.search(DEFENDANT_REGEX, combined, flags=re.IGNORECASE)
    if match:
        return normalize_unicode(match.group(1)).strip()
    return "UNKNOWN DEFENDANT"


def read_registry_number_from_pdf(pdf_path: Path) -> str:
    try:
        reader = PdfReader(str(pdf_path))
    except Exception:
        return ""
    fields = reader.get_fields() or {}
    if "cfn" in fields and fields["cfn"].get("/V"):
        return str(fields["cfn"]["/V"]).strip()
    for key, value in fields.items():
        haystack = (
            (value.get("/T") or "")
            + " "
            + (value.get("/TU") or "")
            + " "
            + (key or "")
        ).lower()
        if any(token in haystack for token in ("cfn", "file", "registry")):
            val = value.get("/V")
            if val:
                return str(val).strip()
    return ""


# ---------------------------------------------------------------------------
# PDF helpers

def ensure_need_appearances(writer: PdfWriter) -> None:
    try:
        writer._root_object["/AcroForm"].update(
            {NameObject("/NeedAppearances"): BooleanObject(True)}
        )
    except KeyError:
        pass


def write_field_all_pages(writer: PdfWriter, field_key: str, value: str) -> None:
    if not field_key or value is None:
        return
    for page in writer.pages:
        if "/Annots" in page:
            writer.update_page_form_field_values(page, {field_key: value})


def find_defendant_field_keys(reader: PdfReader) -> list[str]:
    fields = reader.get_fields() or {}
    keys: list[str] = []
    if "defsoc" in fields:
        keys.append("defsoc")
    for key, value in fields.items():
        haystack = (
            (value.get("/T") or "")
            + " "
            + (value.get("/TU") or "")
            + " "
            + (key or "")
        ).lower()
        if "defendant" in haystack and key not in keys:
            keys.append(key)
    return keys


def find_claim_against_field(reader: PdfReader) -> Optional[str]:
    fields = reader.get_fields() or {}
    if "claimagainst" in fields:
        return "claimagainst"
    for key, value in fields.items():
        haystack = (
            (value.get("/T") or "")
            + " "
            + (value.get("/TU") or "")
            + " "
            + (key or "")
        ).lower()
        if "claim against" in haystack or "claimagainst" in haystack:
            return key
    return None


# ---------------------------------------------------------------------------
# Core workflow

def fill_dismissal_form(
    claim_docx: Path,
    notice_pdf: Path,
    template_pdf: Path,
    output_dir: Path,
) -> tuple[Path, dict[str, str]]:
    defendant_name = extract_defendant_from_claim(claim_docx)
    registry_number = read_registry_number_from_pdf(notice_pdf) if notice_pdf.exists() else ""

    reader = PdfReader(str(template_pdf))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)

    defendant_keys = find_defendant_field_keys(reader)
    claim_against_key = find_claim_against_field(reader)
    fields = reader.get_fields() or {}
    cfn_key = "cfn" if "cfn" in fields else None

    for key in defendant_keys:
        write_field_all_pages(writer, key, defendant_name)

    if claim_against_key:
        write_field_all_pages(writer, claim_against_key, defendant_name)

    if registry_number and cfn_key:
        write_field_all_pages(writer, cfn_key, registry_number)

    ensure_need_appearances(writer)
    output_dir.mkdir(parents=True, exist_ok=True)
    out_name = f"BC Dismissals - {safe_filename(defendant_name)}.pdf"
    output_path = output_dir / out_name
    with open(output_path, "wb") as file_obj:
        writer.write(file_obj)

    summary = {
        "defendant": defendant_name,
        "registry_number": registry_number or "(blank)",
        "defendant_fields": ", ".join(defendant_keys) or "(none found)",
        "claim_against_field": claim_against_key or "(not found)",
        "cfn_field": cfn_key or "(not found)",
    }
    return output_path, summary


# ---------------------------------------------------------------------------
# CLI helpers

def auto_pick_file(directory: Path, patterns: Iterable[str]) -> Optional[Path]:
    if not directory.exists():
        return None
    for pattern in patterns:
        matches = sorted(
            directory.glob(pattern),
            key=lambda candidate: candidate.stat().st_mtime,
            reverse=True,
        )
        if matches:
            return matches[0]
    return None


def resolve_inputs_from_args(args: argparse.Namespace) -> tuple[Path, Path, Path, Path]:
    claim_docx = Path(args.claim_docx).expanduser().resolve() if args.claim_docx else None
    notice_pdf = Path(args.notice_pdf).expanduser().resolve() if args.notice_pdf else None
    template_pdf = (
        Path(args.template_pdf).expanduser().resolve() if args.template_pdf else None
    )
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else None

    if claim_docx is None:
        claim_docx = auto_pick_file(DEFAULT_INPUT_DIR, ("*Schedule*.docx", "*.docx"))
    if notice_pdf is None:
        notice_pdf = auto_pick_file(DEFAULT_INPUT_DIR, ("*Notice*.pdf", "*.pdf"))
    if template_pdf is None:
        template_pdf = ASSET_TEMPLATE
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR

    required = [
        ("Schedule A DOCX", claim_docx),
        ("Notice of Claim PDF", notice_pdf),
        ("Dismissal template PDF", template_pdf),
    ]
    for label, path in required:
        if path is None or not path.exists():
            raise FileNotFoundError(f"Unable to locate {label}. Checked path: {path}")

    return claim_docx, notice_pdf, template_pdf, output_dir


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Populate the BC Small Claims Notice of Withdrawal form."
    )
    parser.add_argument("--claim-docx", help="Path to the Schedule A / claim DOCX.")
    parser.add_argument("--notice-pdf", help="Path to the Notice of Claim PDF.")
    parser.add_argument(
        "--template-pdf",
        help="Path to the Notice of Withdrawal PDF template.",
    )
    parser.add_argument(
        "--output-dir",
        help="Directory to store the populated PDF (default: ./output).",
    )
    return parser.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> None:
    args = parse_args(argv)
    claim_docx, notice_pdf, template_pdf, output_dir = resolve_inputs_from_args(args)

    output_path, summary = fill_dismissal_form(
        claim_docx=claim_docx,
        notice_pdf=notice_pdf,
        template_pdf=template_pdf,
        output_dir=output_dir,
    )

    print(f"Saved dismissal to: {output_path}")
    print("Summary:")
    for key, value in summary.items():
        print(f"  {key.replace('_', ' ').title()}: {value}")


if __name__ == "__main__":
    main()
